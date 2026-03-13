import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode

def create_instructions_doc():
    doc_path = r"c:\Users\Ruken Kuvvetli\Desktop\Word Weaver T1 - Bayram\Word_Weaver_Instructions.docx"
    qr_path = r"c:\Users\Ruken Kuvvetli\Desktop\Word Weaver T1 - Bayram\images\qr_code.png"
    game_link = "https://meb.ai/UK4BU2z"

    # Generate QR code
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(game_link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(qr_path)

    # Create Document
    doc = Document()
    
    # Title
    title = doc.add_heading('How to Play Word Weaver: Ramadan Eid', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Welcome to Word Weaver! This is a fun and easy word game about important values in Ramadan Eid.')

    # Rules Heading
    doc.add_heading('Game Rules', level=1)
    
    rules = [
        "1. Read the Clue: Look at the clue box on the screen. It tells you the meaning of the hidden word.",
        "2. Find the Letters: The missing letters are in the colorful flying candies!",
        "3. Click the Right Letters: Click on the candies in the right order to spell the word.",
        "4. Check Your Answer: If you click the right letter, the box turns green. If you click the wrong letter, it turns red and you lose a life (❤️).",
        "5. Be Careful: You only have 3 lives! If you lose all 3 lives, the game is over.",
        "6. Win the Game: Spell all 7 words correctly to win the game!"
    ]
    
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')

    # Buttons Heading
    doc.add_heading('Game Buttons', level=1)
    
    buttons = [
        "Shuffle: Mixes the flying letters so you can see them better.",
        "Skip: Jump to a new word if you do not know the answer. But try your best first!",
        "Music: Turn the background music on or off.",
        "Bright / Dark (☀️ / 🌙): Change the background to a sunny day or a beautiful night!"
    ]

    for button in buttons:
        doc.add_paragraph(button, style='List Bullet')

    # Closing
    p = doc.add_paragraph()
    p.add_run('\nEnjoy playing and sharing these beautiful values with your friends!').bold = True

    # Link and QR
    doc.add_heading('Play on your Phone or Tablet!', level=2)
    doc.add_paragraph(f'Game Link: {game_link}')
    
    # Add QR code image
    doc.add_picture(qr_path, width=Inches(2.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save Document
    doc.save(doc_path)
    print(f"Successfully created {doc_path}")

if __name__ == '__main__':
    create_instructions_doc()
